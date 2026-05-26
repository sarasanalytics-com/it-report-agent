#!/usr/bin/env python3
"""
Convert the markdown report + metrics.json + charts into a styled Word doc.

Usage:
    python generate-docx.py <markdown_file> <output_docx> [--metrics PATH] [--charts DIR]

Looks for:
  output/metrics.json  — structured KPIs (see generate-report.py)
  output/charts/*.png  — pre-rendered visuals (see generate-charts.py)
"""

import json
import re
import sys
import pathlib

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor, Cm

# ── Brand palette (kept in sync with generate-charts.py) ──────────────────
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
INK = RGBColor(0x21, 0x31, 0x4D)
MUTED = RGBColor(0x6B, 0x7A, 0x99)
GREEN = RGBColor(0x43, 0xA0, 0x47)
AMBER = RGBColor(0xE9, 0xB9, 0x49)
RED = RGBColor(0xE6, 0x39, 0x46)
LIGHT_BG = "F4F6FA"


# ── XML helpers ───────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background shading."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _no_borders(table) -> None:
    """Remove all borders from a table (for cover/KPI panels)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    tblPr.append(borders)


def _status_color(icon: str) -> RGBColor:
    return {"🟢": GREEN, "🟡": AMBER, "🔴": RED}.get(icon, MUTED)


def _status_hex(icon: str) -> str:
    return {"🟢": "43A047", "🟡": "E9B949", "🔴": "E63946"}.get(icon, "6B7A99")


# ── Document sections ─────────────────────────────────────────────────────

def add_cover_page(doc: Document, metrics: dict, charts_dir: pathlib.Path) -> None:
    """Hero banner + KPI strip rendered as an image at the top."""
    kpi_img = charts_dir / "kpi-strip.png"
    if kpi_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(kpi_img), width=Inches(7.0))
    else:
        # Fallback: text-only header
        title = doc.add_paragraph()
        run = title.add_run("IT Operations Report")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = NAVY
        sub = doc.add_paragraph()
        sub_run = sub.add_run(f"{metrics.get('date', '')} · {metrics.get('overall_status', '')}")
        sub_run.font.size = Pt(11)
        sub_run.font.color.rgb = MUTED


def add_executive_summary(doc: Document, metrics: dict) -> None:
    """One-paragraph summary + risks callout."""
    doc.add_paragraph()  # spacer
    h = doc.add_paragraph()
    hr = h.add_run("Executive Summary")
    hr.font.size = Pt(14)
    hr.font.bold = True
    hr.font.color.rgb = NAVY

    kpis = metrics["kpis"]
    risks = metrics.get("risks", [])
    status = metrics.get("overall_status", "")

    runway_note = ""
    if kpis.get("runway_weeks") is not None:
        runway_note = f" Procurement runway sits at ~{kpis['runway_weeks']} weeks at current pace."
    spend_note = ""
    if kpis.get("laptop_spend_pct_of_budget") is not None:
        spend_note = f" Month-to-date laptop spend is {kpis['laptop_spend_pct_of_budget']:.0f}% of monthly budget."

    summary = (
        f"Overall status: {status}. "
        f"{kpis['total_laptops']} laptops in fleet ({kpis['total_assigned']} assigned, "
        f"{kpis['stock_ready']} ready in stock, {kpis['stock_backup']} backup). "
        f"{kpis['aging_total']} laptops past 3.5 years ({kpis['aging_critical']} critical). "
        f"{kpis['joiners_next_7']} joiners next 7 days, {kpis['joiners_next_30']} next 30 days."
        f"{runway_note}{spend_note}"
    )
    p = doc.add_paragraph(summary)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = INK

    if risks:
        callout = doc.add_table(rows=1, cols=1)
        callout.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = callout.rows[0].cells[0]
        _set_cell_bg(cell, "FFF4D6")
        cell.width = Inches(7.0)
        cell.paragraphs[0].clear()
        title_run = cell.paragraphs[0].add_run("⚠️ Risks & Actions")
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.color.rgb = INK
        for r in risks:
            line = cell.add_paragraph()
            line.paragraph_format.left_indent = Inches(0.1)
            for part in re.split(r"(\*[^*]+\*)", r):
                run = line.add_run(part.strip("*"))
                run.font.size = Pt(10)
                run.font.color.rgb = INK
                if part.startswith("*") and part.endswith("*"):
                    run.bold = True


def add_chart_row(doc: Document, charts_dir: pathlib.Path) -> None:
    """Aging donut + Stock vs Demand side by side."""
    aging = charts_dir / "aging.png"
    stock = charts_dir / "stock.png"
    if not (aging.exists() or stock.exists()):
        return

    doc.add_paragraph()  # spacer
    table = doc.add_table(rows=1, cols=2)
    _no_borders(table)
    table.autofit = False
    for c in table.columns:
        c.width = Inches(3.5)

    if aging.exists():
        cell = table.rows[0].cells[0]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run().add_picture(str(aging), width=Inches(3.4))
    if stock.exists():
        cell = table.rows[0].cells[1]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run().add_picture(str(stock), width=Inches(3.4))


def add_runway_chart(doc: Document, charts_dir: pathlib.Path) -> None:
    runway = charts_dir / "runway.png"
    if runway.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(runway), width=Inches(6.5))


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Saras Analytics · IT Operations · Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def set_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)


# ── Markdown → Word (kept from prior version, with styling tweaks) ────────

def add_table_from_md(doc: Document, header_line: str, rows: list) -> None:
    def parse_cells(line: str) -> list:
        return [c.strip() for c in line.strip().strip("|").split("|")]

    headers = parse_cells(header_line)
    num_cols = len(headers)

    table = doc.add_table(rows=1, cols=num_cols, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row — navy background + white text
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_bg(cell, "1F3A5F")
        cell.text = ""
        run = cell.paragraphs[0].add_run(h.replace("**", ""))
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for idx, row_line in enumerate(rows):
        cells = parse_cells(row_line)
        row = table.add_row()
        for i, val in enumerate(cells[:num_cols]):
            cell = row.cells[i]
            cell.text = ""
            # Zebra striping
            if idx % 2 == 1:
                _set_cell_bg(cell, LIGHT_BG)
            # Status pill coloring on the right-most or first cell containing emoji
            stripped_val = val.replace("**", "")
            run = cell.paragraphs[0].add_run(stripped_val)
            run.font.size = Pt(9.5)
            run.font.color.rgb = INK
            if "**" in val:
                run.bold = True
            # Color status text
            if "🟢" in stripped_val:
                run.font.color.rgb = GREEN
                run.bold = True
            elif "🟡" in stripped_val:
                run.font.color.rgb = AMBER
                run.bold = True
            elif "🔴" in stripped_val:
                run.font.color.rgb = RED
                run.bold = True


def _styled_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = NAVY
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = INK


def md_to_docx(md_text: str, output_path: pathlib.Path,
               metrics_path: pathlib.Path, charts_dir: pathlib.Path) -> None:
    doc = Document()
    set_page_margins(doc)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Cover (hero KPI strip) ──
    metrics = {}
    if metrics_path.exists():
        try:
            metrics = json.loads(metrics_path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            pass

    if metrics:
        add_cover_page(doc, metrics, charts_dir)
        add_executive_summary(doc, metrics)
        add_chart_row(doc, charts_dir)
        add_runway_chart(doc, charts_dir)
        doc.add_paragraph()

    # ── Body from markdown ──
    lines = md_text.split("\n")
    i = 0
    # Skip the original H1 since the cover already shows the title
    if metrics and lines and lines[0].startswith("# "):
        i = 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            i += 1
            continue

        if stripped.startswith("# "):
            _styled_heading(doc, stripped[2:], 1)
            i += 1
            continue
        if stripped.startswith("## "):
            _styled_heading(doc, stripped[3:], 2)
            i += 1
            continue
        if stripped.startswith("### "):
            _styled_heading(doc, stripped[4:], 3)
            i += 1
            continue

        # Skip the inline Health Check / Risks / Headline Metrics sections that
        # are already shown beautifully in the cover/exec summary.
        if metrics and stripped in (
            "🚦 Health Check", "⚠️ Risks & Actions", "📊 Headline Metrics",
        ):
            # also skip until next "## " heading
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("## "):
                i += 1
            continue

        # Tables
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\|[-| :]+\|$", lines[i + 1].strip()):
            header_line = stripped
            i += 2
            table_rows = []
            while i < len(lines) and "|" in lines[i].strip() and lines[i].strip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            add_table_from_md(doc, header_line, table_rows)
            doc.add_paragraph("")
            continue

        if stripped.startswith("**"):
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            i += 1
            continue

        if stripped.startswith("• ") or stripped.startswith("- "):
            text = stripped[2:]
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue

        if stripped.startswith("  - ") or stripped.startswith("  • "):
            text = stripped[4:]
            doc.add_paragraph(text, style="List Bullet 2")
            i += 1
            continue

        if stripped.startswith("_") and stripped.endswith("_"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("_"))
            run.italic = True
            run.font.size = Pt(8)
            run.font.color.rgb = MUTED
            i += 1
            continue

        doc.add_paragraph(stripped)
        i += 1

    add_footer(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    size_kb = output_path.stat().st_size / 1024
    print(f"Generated: {output_path} ({size_kb:.0f} KB)")


def main() -> None:
    if len(sys.argv) < 3:
        print("Usage: generate-docx.py <markdown_file> <output_docx>", file=sys.stderr)
        sys.exit(1)

    md_path = pathlib.Path(sys.argv[1])
    out_path = pathlib.Path(sys.argv[2])

    if not md_path.exists():
        print(f"File not found: {md_path}", file=sys.stderr)
        sys.exit(1)

    metrics_path = md_path.parent / "metrics.json"
    charts_dir = md_path.parent / "charts"

    content = md_path.read_text(encoding="utf-8").strip()
    if not content:
        print("Markdown file is empty — skipping.", file=sys.stderr)
        sys.exit(1)

    md_to_docx(content, out_path, metrics_path, charts_dir)


if __name__ == "__main__":
    main()
